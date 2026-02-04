import json
import os
from pathlib import Path

from docx import Document as DocxDocument
from openai import OpenAI

from app.config import settings


class ResumeParser:
    """Parses uploaded DOCX resumes into structured JSON data."""

    def __init__(self):
        self.client = OpenAI(api_key=settings.OPENAI_API_KEY)
        self.model = settings.OPENAI_MODEL

    def extract_text(self, docx_path: str) -> str:
        """Extract raw text from a DOCX file."""
        doc = DocxDocument(docx_path)
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())

        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))

        return "\n".join(text_parts)

    def parse(self, docx_path: str) -> dict:
        """Parse a DOCX resume into structured JSON."""
        raw_text = self.extract_text(docx_path)
        return self._structure_with_llm(raw_text)

    def _structure_with_llm(self, raw_text: str) -> dict:
        """Use GPT to extract structured data from resume text."""
        response = self.client.chat.completions.create(
            model=self.model,
            response_format={"type": "json_object"},
            messages=[
                {
                    "role": "system",
                    "content": """You are a resume parser. Extract structured data from the resume text.
Return a JSON object with these exact fields:
{
    "full_name": "string",
    "email": "string",
    "phone": "string or null",
    "linkedin_url": "string or null",
    "city": "string or null",
    "state": "string or null",
    "country": "string or null",
    "summary": "string - professional summary/objective",
    "skills": ["array of skill strings"],
    "experience": [
        {
            "company": "string",
            "title": "string",
            "start_date": "string",
            "end_date": "string or Present",
            "location": "string or null",
            "bullets": ["array of achievement/responsibility strings"]
        }
    ],
    "education": [
        {
            "school": "string",
            "degree": "string",
            "field": "string",
            "graduation_date": "string or null",
            "gpa": "string or null"
        }
    ],
    "certifications": ["array of certification strings"],
    "languages": ["array of language strings"],
    "projects": [
        {
            "name": "string",
            "description": "string",
            "technologies": ["array of tech strings"]
        }
    ]
}

Be thorough and extract ALL information present. If a field is not found, use null or empty array.""",
                },
                {"role": "user", "content": raw_text},
            ],
        )

        return json.loads(response.choices[0].message.content)
