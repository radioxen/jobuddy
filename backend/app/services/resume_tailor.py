import json
import os
from datetime import datetime

from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openai import OpenAI

from app.config import settings


class ResumeTailorService:
    """Tailors resume content for specific job descriptions and generates DOCX."""

    def __init__(self):
        self.client = OpenAI(api_key=settings.OPENAI_API_KEY)
        self.model = settings.OPENAI_MODEL

    def tailor(
        self,
        resume_data: dict,
        job_description: str,
        job_title: str,
        company: str,
        additional_context: str = "",
    ) -> dict:
        """Use GPT to tailor resume content for a specific job.

        Args:
            resume_data: Parsed resume data dictionary
            job_description: The job posting description
            job_title: Target job title
            company: Target company name
            additional_context: Optional RAG-retrieved context (portfolio, projects, etc.)
        """
        # Build the user prompt with optional additional context
        user_content = f"""Original resume data:
{json.dumps(resume_data, indent=2)}

Target job:
Title: {job_title}
Company: {company}
Description: {job_description}"""

        if additional_context:
            user_content += f"""

Additional context from candidate's portfolio/projects (use to enhance relevant sections):
{additional_context}"""

        user_content += "\n\nPlease tailor the resume for this specific job."

        response = self.client.chat.completions.create(
            model=self.model,
            response_format={"type": "json_object"},
            messages=[
                {
                    "role": "system",
                    "content": """You are a professional resume writer. Tailor the candidate's resume for the specific job.

Rules:
1. Adjust the professional summary to align with the job requirements
2. Reorder skills to prioritize the most relevant ones for this job
3. Rephrase experience bullets to highlight relevant achievements using strong action verbs
4. Quantify achievements where possible
5. Keep ALL information truthful — never fabricate experience or skills
6. Optimize for ATS keyword matching
7. Keep the same structure but improve relevance
8. If additional context (portfolio/projects) is provided, incorporate relevant details that strengthen the application

Return a JSON object with the SAME structure as the input resume data, but with tailored content.
The JSON must include: full_name, email, phone, linkedin_url, summary, skills, experience, education, certifications.""",
                },
                {
                    "role": "user",
                    "content": user_content,
                },
            ],
        )

        return json.loads(response.choices[0].message.content)

    def generate_docx(self, tailored_data: dict, output_path: str) -> str:
        """Generate a well-formatted DOCX resume from tailored data."""
        doc = DocxDocument()

        # Set default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)
        font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # --- Header: Name ---
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(tailored_data.get("full_name", ""))
        name_run.font.size = Pt(22)
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        # --- Contact Info ---
        contact_parts = []
        if tailored_data.get("email"):
            contact_parts.append(tailored_data["email"])
        if tailored_data.get("phone"):
            contact_parts.append(tailored_data["phone"])
        if tailored_data.get("linkedin_url"):
            contact_parts.append(tailored_data["linkedin_url"])
        if tailored_data.get("city"):
            loc = tailored_data["city"]
            if tailored_data.get("state"):
                loc += f", {tailored_data['state']}"
            contact_parts.append(loc)

        if contact_parts:
            contact_para = doc.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_run = contact_para.add_run(" | ".join(contact_parts))
            contact_run.font.size = Pt(10)
            contact_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # --- Professional Summary ---
        if tailored_data.get("summary"):
            self._add_section_heading(doc, "PROFESSIONAL SUMMARY")
            summary_para = doc.add_paragraph(tailored_data["summary"])
            summary_para.paragraph_format.space_after = Pt(6)

        # --- Skills ---
        if tailored_data.get("skills"):
            self._add_section_heading(doc, "SKILLS")
            skills_text = " • ".join(tailored_data["skills"])
            skills_para = doc.add_paragraph(skills_text)
            skills_para.paragraph_format.space_after = Pt(6)

        # --- Experience ---
        if tailored_data.get("experience"):
            self._add_section_heading(doc, "PROFESSIONAL EXPERIENCE")
            for exp in tailored_data["experience"]:
                # Job title and company on same line
                title_para = doc.add_paragraph()
                title_run = title_para.add_run(
                    f"{exp.get('title', '')} — {exp.get('company', '')}"
                )
                title_run.font.bold = True
                title_run.font.size = Pt(11)

                # Date and location
                date_parts = []
                if exp.get("start_date"):
                    end = exp.get("end_date", "Present")
                    date_parts.append(f"{exp['start_date']} – {end}")
                if exp.get("location"):
                    date_parts.append(exp["location"])

                if date_parts:
                    date_para = doc.add_paragraph()
                    date_run = date_para.add_run(" | ".join(date_parts))
                    date_run.font.size = Pt(10)
                    date_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                    date_run.font.italic = True

                # Bullets
                for bullet in exp.get("bullets", []):
                    bullet_para = doc.add_paragraph(bullet, style="List Bullet")
                    bullet_para.paragraph_format.space_after = Pt(2)

        # --- Education ---
        if tailored_data.get("education"):
            self._add_section_heading(doc, "EDUCATION")
            for edu in tailored_data["education"]:
                edu_para = doc.add_paragraph()
                degree_text = edu.get("degree", "")
                if edu.get("field"):
                    degree_text += f" in {edu['field']}"
                edu_run = edu_para.add_run(degree_text)
                edu_run.font.bold = True

                school_parts = [edu.get("school", "")]
                if edu.get("graduation_date"):
                    school_parts.append(edu["graduation_date"])
                if edu.get("gpa"):
                    school_parts.append(f"GPA: {edu['gpa']}")

                school_para = doc.add_paragraph(" | ".join(school_parts))
                school_para.paragraph_format.space_after = Pt(4)

        # --- Certifications ---
        if tailored_data.get("certifications"):
            self._add_section_heading(doc, "CERTIFICATIONS")
            for cert in tailored_data["certifications"]:
                doc.add_paragraph(f"• {cert}")

        # --- Projects ---
        if tailored_data.get("projects"):
            self._add_section_heading(doc, "PROJECTS")
            for proj in tailored_data["projects"]:
                proj_para = doc.add_paragraph()
                proj_run = proj_para.add_run(proj.get("name", ""))
                proj_run.font.bold = True
                if proj.get("description"):
                    doc.add_paragraph(proj["description"])
                if proj.get("technologies"):
                    tech_para = doc.add_paragraph(
                        f"Technologies: {', '.join(proj['technologies'])}"
                    )
                    tech_para.paragraph_format.space_after = Pt(4)

        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)

        doc.save(output_path)
        return output_path

    def _add_section_heading(self, doc: DocxDocument, title: str):
        """Add a formatted section heading with a line underneath."""
        heading_para = doc.add_paragraph()
        heading_para.paragraph_format.space_before = Pt(12)
        heading_para.paragraph_format.space_after = Pt(4)
        heading_run = heading_para.add_run(title)
        heading_run.font.size = Pt(12)
        heading_run.font.bold = True
        heading_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        # Add a thin line (border bottom effect via paragraph border)
        from docx.oxml.ns import qn
        pPr = heading_para._p.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pBdr.makeelement(
            qn("w:bottom"),
            {
                qn("w:val"): "single",
                qn("w:sz"): "4",
                qn("w:space"): "1",
                qn("w:color"): "1A1A2E",
            },
        )
        pBdr.append(bottom)
        pPr.append(pBdr)
