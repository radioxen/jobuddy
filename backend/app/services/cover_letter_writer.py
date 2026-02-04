import json
import os

from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openai import OpenAI

from app.config import settings


class CoverLetterWriter:
    """Generates personalized cover letters and outputs them as DOCX."""

    def __init__(self):
        self.client = OpenAI(api_key=settings.OPENAI_API_KEY)
        self.model = settings.OPENAI_MODEL

    def write(
        self,
        resume_data: dict,
        job_title: str,
        company: str,
        job_description: str,
    ) -> str:
        """Generate a cover letter using GPT. Returns the letter text."""
        response = self.client.chat.completions.create(
            model=self.model,
            messages=[
                {
                    "role": "system",
                    "content": """You are a professional cover letter writer. Write a compelling,
personalized cover letter that connects the candidate's experience to the specific job requirements.

Rules:
1. Open with a compelling hook — not "I am writing to apply for..."
2. Connect 2-3 specific experiences/achievements from the resume to job requirements
3. Show knowledge of the company when possible
4. Keep it under 400 words
5. Close with enthusiasm and a call to action
6. Be professional but authentic — avoid clichés and generic phrases
7. Use a conversational yet professional tone

Format: Plain text paragraphs. Do NOT include the address header or date — just the letter body.
Start directly with the salutation (Dear Hiring Manager,) and end with the sign-off.""",
                },
                {
                    "role": "user",
                    "content": f"""Candidate resume data:
{json.dumps(resume_data, indent=2)}

Target job:
Title: {job_title}
Company: {company}
Description: {job_description}

Write a cover letter for this specific job.""",
                },
            ],
        )

        return response.choices[0].message.content

    def generate_docx(
        self,
        letter_text: str,
        candidate_name: str,
        candidate_email: str,
        candidate_phone: str,
        output_path: str,
    ) -> str:
        """Generate a well-formatted cover letter DOCX."""
        doc = DocxDocument()

        # Set default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)
        font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # --- Header: Candidate Name ---
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        name_run = name_para.add_run(candidate_name)
        name_run.font.size = Pt(16)
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        # Contact info
        contact_parts = []
        if candidate_email:
            contact_parts.append(candidate_email)
        if candidate_phone:
            contact_parts.append(candidate_phone)
        if contact_parts:
            contact_para = doc.add_paragraph(" | ".join(contact_parts))
            contact_para.paragraph_format.space_after = Pt(12)

        # Add a separator line
        separator = doc.add_paragraph()
        separator.paragraph_format.space_after = Pt(12)
        from docx.oxml.ns import qn

        pPr = separator._p.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pBdr.makeelement(
            qn("w:bottom"),
            {
                qn("w:val"): "single",
                qn("w:sz"): "4",
                qn("w:space"): "1",
                qn("w:color"): "1A1A2E",
            },
        )
        pBdr.append(bottom)
        pPr.append(pBdr)

        # --- Letter body ---
        paragraphs = letter_text.strip().split("\n\n")
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue
            para = doc.add_paragraph(para_text)
            para.paragraph_format.space_after = Pt(8)
            para.paragraph_format.line_spacing = 1.15

        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

        doc.save(output_path)
        return output_path
