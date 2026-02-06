"""
RAG (Retrieval Augmented Generation) Service

Loads documents (resume, portfolio, etc.) and provides relevant context
to AI agents when tailoring resumes and writing cover letters.
"""

import os
import json
from pathlib import Path
from typing import Optional

from docx import Document as DocxDocument
from openai import OpenAI

from app.config import settings


class DocumentChunk:
    """Represents a chunk of a document with metadata."""

    def __init__(self, content: str, source: str, chunk_type: str, metadata: dict = None):
        self.content = content
        self.source = source  # filename
        self.chunk_type = chunk_type  # "resume", "portfolio", "project", etc.
        self.metadata = metadata or {}

    def to_dict(self) -> dict:
        return {
            "content": self.content,
            "source": self.source,
            "chunk_type": self.chunk_type,
            "metadata": self.metadata,
        }


class RAGService:
    """
    Retrieval Augmented Generation service.

    Loads documents from the document folder, chunks them intelligently,
    and provides relevant context for AI agents.
    """

    def __init__(self, document_dir: str = None):
        self.document_dir = document_dir or os.path.join(
            os.path.dirname(settings.BASE_DIR), "documents"
        )
        self.client = OpenAI(api_key=settings.OPENAI_API_KEY)
        self.model = settings.OPENAI_MODEL
        self.chunks: list[DocumentChunk] = []
        self._loaded = False

    def load_documents(self) -> int:
        """Load and chunk all documents from the document directory."""
        if not os.path.exists(self.document_dir):
            print(f"Document directory not found: {self.document_dir}")
            return 0

        self.chunks = []
        count = 0

        for filename in os.listdir(self.document_dir):
            filepath = os.path.join(self.document_dir, filename)
            if filename.endswith(".docx"):
                try:
                    doc_chunks = self._load_docx(filepath, filename)
                    self.chunks.extend(doc_chunks)
                    count += 1
                    print(f"Loaded {filename}: {len(doc_chunks)} chunks")
                except Exception as e:
                    print(f"Error loading {filename}: {e}")
            elif filename.endswith(".pdf"):
                try:
                    doc_chunks = self._load_pdf(filepath, filename)
                    self.chunks.extend(doc_chunks)
                    count += 1
                    print(f"Loaded {filename}: {len(doc_chunks)} chunks")
                except Exception as e:
                    print(f"Error loading {filename}: {e}")
            elif filename.endswith(".txt"):
                try:
                    doc_chunks = self._load_txt(filepath, filename)
                    self.chunks.extend(doc_chunks)
                    count += 1
                except Exception as e:
                    print(f"Error loading {filename}: {e}")

        self._loaded = True
        print(f"RAG: Loaded {count} documents, {len(self.chunks)} total chunks")
        return count

    def _load_docx(self, filepath: str, filename: str) -> list[DocumentChunk]:
        """Load and chunk a DOCX file."""
        doc = DocxDocument(filepath)
        chunks = []

        # Determine document type from filename
        filename_lower = filename.lower()
        if "resume" in filename_lower or "cv" in filename_lower:
            chunk_type = "resume"
        elif "portfolio" in filename_lower:
            chunk_type = "portfolio"
        elif "cover" in filename_lower:
            chunk_type = "cover_letter_sample"
        else:
            chunk_type = "document"

        # Extract full text first
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())

        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))

        # Create one main chunk with full content
        main_content = "\n".join(full_text)
        if main_content:
            chunks.append(DocumentChunk(
                content=main_content,
                source=filename,
                chunk_type=chunk_type,
                metadata={"full_document": True}
            ))

        # Also create section-based chunks for better retrieval
        current_section = []
        current_heading = "Introduction"

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Detect headings (usually bold or larger font, or ALL CAPS)
            is_heading = (
                para.style.name.startswith("Heading") or
                text.isupper() or
                (len(text) < 50 and text.endswith(":"))
            )

            if is_heading and current_section:
                # Save previous section
                section_text = "\n".join(current_section)
                if len(section_text) > 100:
                    chunks.append(DocumentChunk(
                        content=section_text,
                        source=filename,
                        chunk_type=chunk_type,
                        metadata={"section": current_heading}
                    ))
                current_section = []
                current_heading = text
            else:
                current_section.append(text)

        # Don't forget last section
        if current_section:
            section_text = "\n".join(current_section)
            if len(section_text) > 100:
                chunks.append(DocumentChunk(
                    content=section_text,
                    source=filename,
                    chunk_type=chunk_type,
                    metadata={"section": current_heading}
                ))

        return chunks

    def _load_pdf(self, filepath: str, filename: str) -> list[DocumentChunk]:
        """Load and chunk a PDF file using pdfplumber."""
        try:
            import pdfplumber
        except ImportError:
            print("pdfplumber not installed, skipping PDF")
            return []

        chunks = []
        chunk_type = "document"
        if "resume" in filename.lower():
            chunk_type = "resume"
        elif "portfolio" in filename.lower():
            chunk_type = "portfolio"

        with pdfplumber.open(filepath) as pdf:
            full_text = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    full_text.append(text)

            if full_text:
                chunks.append(DocumentChunk(
                    content="\n".join(full_text),
                    source=filename,
                    chunk_type=chunk_type,
                    metadata={"pages": len(pdf.pages)}
                ))

        return chunks

    def _load_txt(self, filepath: str, filename: str) -> list[DocumentChunk]:
        """Load a text file."""
        with open(filepath, "r", encoding="utf-8") as f:
            content = f.read()

        return [DocumentChunk(
            content=content,
            source=filename,
            chunk_type="document",
            metadata={}
        )]

    def get_context_for_job(self, job_description: str, job_title: str, company: str) -> str:
        """
        Get relevant context from loaded documents for a specific job.
        Uses GPT to select and summarize the most relevant information.
        """
        if not self._loaded:
            self.load_documents()

        if not self.chunks:
            return ""

        # Gather all document content
        resume_content = ""
        portfolio_content = ""
        other_content = ""

        for chunk in self.chunks:
            if chunk.chunk_type == "resume":
                resume_content += f"\n{chunk.content}"
            elif chunk.chunk_type == "portfolio":
                portfolio_content += f"\n{chunk.content}"
            else:
                other_content += f"\n{chunk.content}"

        # Build context prompt
        context_parts = []
        if resume_content:
            context_parts.append(f"=== CANDIDATE RESUME ===\n{resume_content[:8000]}")
        if portfolio_content:
            context_parts.append(f"=== COMPANY PORTFOLIO (Candidate's Company) ===\n{portfolio_content[:8000]}")
        if other_content:
            context_parts.append(f"=== ADDITIONAL DOCUMENTS ===\n{other_content[:4000]}")

        full_context = "\n\n".join(context_parts)

        # Use GPT to extract the most relevant information
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {
                        "role": "system",
                        "content": """You are a context extraction assistant. Given documents about a candidate
and a job they're applying to, extract and summarize the MOST RELEVANT information that would help
tailor their application. Focus on:
1. Relevant skills and experience from the resume
2. Relevant projects/case studies from the portfolio that match the job
3. Achievements that align with the job requirements
4. Any unique selling points

Keep your summary concise but comprehensive (max 2000 words)."""
                    },
                    {
                        "role": "user",
                        "content": f"""Job being applied to:
Title: {job_title}
Company: {company}
Description: {job_description[:3000]}

Candidate documents:
{full_context}

Extract the most relevant information for this specific job application."""
                    }
                ],
                max_tokens=2500,
            )
            return response.choices[0].message.content
        except Exception as e:
            print(f"RAG context extraction error: {e}")
            # Fallback: return raw content
            return full_context[:5000]

    def get_resume_data(self) -> Optional[dict]:
        """Extract structured resume data using GPT."""
        if not self._loaded:
            self.load_documents()

        resume_chunks = [c for c in self.chunks if c.chunk_type == "resume"]
        if not resume_chunks:
            return None

        resume_text = "\n".join(c.content for c in resume_chunks)

        try:
            response = self.client.chat.completions.create(
                model=self.model,
                response_format={"type": "json_object"},
                messages=[
                    {
                        "role": "system",
                        "content": """Extract structured resume data. Return JSON with:
{
    "full_name": "string",
    "email": "string",
    "phone": "string or null",
    "linkedin_url": "string or null",
    "city": "string or null",
    "summary": "professional summary",
    "skills": ["skill1", "skill2"],
    "experience": [{"company": "", "title": "", "start_date": "", "end_date": "", "bullets": []}],
    "education": [{"school": "", "degree": "", "field": "", "graduation_date": ""}],
    "certifications": [],
    "projects": [{"name": "", "description": "", "technologies": []}]
}"""
                    },
                    {"role": "user", "content": resume_text[:10000]}
                ],
            )
            return json.loads(response.choices[0].message.content)
        except Exception as e:
            print(f"Resume parsing error: {e}")
            return None

    def get_portfolio_highlights(self) -> list[dict]:
        """Extract key projects/case studies from portfolio."""
        if not self._loaded:
            self.load_documents()

        portfolio_chunks = [c for c in self.chunks if c.chunk_type == "portfolio"]
        if not portfolio_chunks:
            return []

        portfolio_text = "\n".join(c.content for c in portfolio_chunks)

        try:
            response = self.client.chat.completions.create(
                model=self.model,
                response_format={"type": "json_object"},
                messages=[
                    {
                        "role": "system",
                        "content": """Extract key projects/case studies from this company portfolio.
Return JSON: {"projects": [{"name": "", "client": "", "description": "", "technologies": [], "results": "", "industry": ""}]}"""
                    },
                    {"role": "user", "content": portfolio_text[:10000]}
                ],
            )
            data = json.loads(response.choices[0].message.content)
            return data.get("projects", [])
        except Exception as e:
            print(f"Portfolio parsing error: {e}")
            return []


# Singleton instance
_rag_service: Optional[RAGService] = None


def get_rag_service() -> RAGService:
    global _rag_service
    if _rag_service is None:
        _rag_service = RAGService()
    return _rag_service
